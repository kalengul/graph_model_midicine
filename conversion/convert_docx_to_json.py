"""Модуль конвертирования docx-файлов в json-файлов."""
import json
import os
from pathlib import Path
from typing import Final


from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


DIR_LEAFLET: Final[str] = 'вкладыши'
DIR_DOCX_INSTRUCTIONS: Final[str] = 'инструкции ГРЛС в docx'
DIR_JSON_INSTRUCTIONS: Final[str] = 'инструкции ГРЛС в json'
LEAFLET_INSERT: Final[str] = 'листок-вкладыш'
KEY_SIDE_EFFECTS: Final[str] = 'Побочное действие'
SYSTEM: Final[str] = 'Систем'
ORGAN: Final[str] = 'орган'


def clear_text(text: str) -> str:
    """Функция для отчистки текста от пробельных символов."""
    text = text.replace('\n', ' ')
    text = text.replace('\t', ' ')
    text = text.replace('  ', ' ')
    return text


def get_table_content(table) -> str:
    """Функция для получения текста из таблицы."""
    table_content = ""
    for row in table.rows:
        for cell in row.cells:
            # Добавляем текст каждой ячейки с пробелом
            table_content += cell.text.strip() + " "
            table_content = clear_text(table_content)
    return table_content


def docx2json(docx_file: str, json_file: str) -> bool:
    """Функция конвертирования docx-файлов в json-файлы."""
    doc = Document(docx_file)
    data = dict()
    key = ''
    # print(doc.paragraphs[1].text.strip().lower())
    if LEAFLET_INSERT not in doc.paragraphs[1].text.strip().lower():
        for paragraph in doc.paragraphs:
            if (paragraph.alignment == WD_ALIGN_PARAGRAPH.RIGHT
                and paragraph.text.strip().isdigit()):
                continue
            # print('paragraph.style =', paragraph.style)
            # print('paragraph.text =', paragraph.text)
            for run in paragraph.runs:
                # print('текст =', run.text)
                if (run.text == '' or run.text.isspace()):
                    # print('следующая итерация')
                    continue
                if (run.bold and (not run.italic)
                    and (not run.underline)
                    and ('•' not in run.text)):
                    key = run.text.strip().replace(':', '')
                    data[key] = ''
                    # print('текст ключа =', run.text)
                elif key:
                    data[key] += run.text.strip() + " "
                    # print('текст содержания =', run.text)
                else:
                    # print('текст никуда не попал =', run.text)
                    continue
        # Обработка таблиц
        for table in doc.tables:
            first_row = table.rows[0]
            cell = first_row.cells[0]
            if ORGAN.lower() in cell.text and SYSTEM.lower() in cell.text:
                side_effects_data = ''
                for row in table.rows[2:]:
                    for cell in row.cells:
                        if cell.text.strip() == '-':
                            continue
                        table_content = ''
                        if cell.tables:
                            for cell_table in cell.tables:
                                table_content += get_table_content(
                                    cell_table)
                        cleared_text = table_content + cell.text.strip()
                        cleared_text = clear_text(cleared_text)
                        side_effects_data += cleared_text.strip() + ' '
                data[KEY_SIDE_EFFECTS] += side_effects_data.strip()
        with open(json_file, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=4)
        return True
    else:
        return False


def main() -> None:
    """Тело программы."""
    files: list[str] = os.listdir(DIR_DOCX_INSTRUCTIONS)
    if not os.path.exists(DIR_JSON_INSTRUCTIONS):
        os.mkdir(DIR_JSON_INSTRUCTIONS)
    # docx_file = 'Апиксабан.docx'
    # json_file = 'Апиксабан.json'
    for docx_file in files:
        json_file = Path(docx_file).stem + '.json'
        docx_file = os.path.join(DIR_DOCX_INSTRUCTIONS, docx_file)
        json_file = os.path.join(DIR_JSON_INSTRUCTIONS, json_file)
        if docx2json(docx_file, json_file):
            print(f'Файл {docx_file} успешно конвертирован в {json_file}')
        else:
            print('Файл {docx_file} является не инструкцией,'
                  ' а листом-вкладышем!')


if __name__ == '__main__':
    main()
