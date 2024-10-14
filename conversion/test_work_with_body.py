"""Модуль конвертирования docx-файлов в json-файлов."""
import json
import os
from pathlib import Path
from typing import (Final,
                    Union)
import string
import re


from docx import Document
from docx.oxml import CT_P, CT_Tbl
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.enum.text import WD_ALIGN_PARAGRAPH


LOG_FILE_NAME = 'лог для программы создания инструкций в формате json.txt'
log_file = open(LOG_FILE_NAME, 'w', encoding='utf-8')
DIR_LEAFLET: Final[str] = 'вкладыши'
DIR_DOCX_INSTRUCTIONS: Final[str] = 'инструкции ГРЛС в docx'
DIR_JSON_INSTRUCTIONS: Final[str] = 'инструкции ГРЛС в json'
LEAFLET_INSERT: Final[str] = 'листок-вкладыш'
KEY_SIDE_EFFECTS: Final[str] = 'Побочное действие'
SYSTEM: Final[str] = 'Систем'
ORGAN: Final[str] = 'орган'
SIDEEFFECT: Final[str] = 'Побочное действие'
ADVERSE_REACTIONS: Final[str] = 'Нежелательные реакции'
HEADLINES: Final[list[str]] = [
    'Фармакодинамика',
    'Фармакокинетика',
]
ABBREVIATIONS: Final[list[str]] = [
    'ООО',
    'ФТГ',
    'АО',
    'ЗАО',
    'ОАО',
]
ALLOWED: Final[list[str]] = [
    'КОД АТХ',
]
FORBIDDEN_WORDS: Final[list[str]] = [
    'Либо'
]
FORBIDDEN_SYMBOL_R: Final[str] = '®'
SHY: Final[str] = '\u00AD'
HYPHEN_SIGN: Final[str] = '-'
with open('список частот.txt', 'r', encoding='utf-8') as file:
    frequency_list: list[str] = list()
    frequency: str = file.readline()
    while frequency:
        frequency_list.append(frequency.strip('\n'))
        frequency = file.readline()
with open('список систем органов.txt', 'r', encoding='utf-8') as file:
    organ_systems: list[str] = list()
    organ_system: str = file.readline()
    while organ_system:
        organ_systems.append(organ_system.strip('\n'))
        organ_system = file.readline()
with open('названия колонок ПД.txt', 'r', encoding='utf-8') as file:
    column_names: list[str] = list()
    column_name: str = file.readline()
    while column_name:
        column_names.append(column_name.strip('\n'))
        column_name = file.readline()
frequency_list = sorted(frequency_list, key=len, reverse=True)
organ_systems = sorted(organ_systems, key=len, reverse=True)
column_names = sorted(column_names, key=len, reverse=True)
print('frequency_list =', column_names)


def availability(row: str, items: list[str]) -> bool:
    """Функция проверки наличия одно из элементов списка items в строке row."""
    for item in items:
        if item in row:
            return True
    return False


def find_index(row: str, items: list[str]) -> int:
    """Функция проверки наличия одно из элементов списка items в строке row."""
    failure = -1
    for item in items:
        if item in row:
            return row.find(item)
    return failure


def clear_text(text: str) -> str:
    """Функция для отчистки текста от пробельных символов."""
    text = text.replace('\n', ' ')
    text = text.replace('\t', ' ')
    text = text.replace('  ', ' ')
    return text


def get_table_content(table) -> str:
    """Функция для получения текста из таблицы."""
    table_content = ""
    for row in table.rows:
        for cell in row.cells:
            # Добавляем текст каждой ячейки с пробелом
            table_content += cell.text.strip() + " "
            table_content = clear_text(table_content)
    return table_content


def delete_empty_items(dictionary: dict[str, str]) -> dict[str, str]:
    """Функция для удаления пустых элементов из словаря."""
    keys = list(dictionary.keys())
    for key in keys:
        if dictionary[key] == '' or dictionary[key].isspace():
            del dictionary[key]
    return dictionary


def is_first_letter_uppercase(row: str) -> bool:
    """Функция, проверяющая, что первая буква в строке - заглавная."""
    if len(row) == 0:  # Проверяем, что строка не пустая
        return False
    return row[0].isupper()


def is_first_letter_lowercase(row: str) -> bool:
    """Функция, проверяющая, что первая буква в строке - строчная."""
    if len(row) == 0:  # Проверяем, что строка не пустая
        return False
    return row[0].islower()


def removing_extra_spaces(text: str) -> str:
    """
    Функция удаления лишних пробелов.

    Если будет в текст два и более пробельных символа,
    идущих подрад, они будут заменены одинарных пробелом.
    """
    return re.sub(r'\s{2,}', ' ', text)


def removing_spaces_after_point(text: str) -> str:
    """Функция удаления пробельных символов перед знаками препинания."""
    return re.sub(r'\s+([.,!?;:])', r'\1', text)


def removing_extra_spaces_value(dictionary: dict[str, str]) -> dict[str, str]:
    """Функция удаления лишних пробельных символов из значений словаря."""
    for key, value in dictionary.items():
        text = removing_extra_spaces(value)
        text = removing_spaces_after_point(text)
        dictionary[key] = text.strip()
    return dictionary


def removing_forbidden_symbol(dictionary: dict[str, str],
                              forbidden_symbol: str) -> dict[str, str]:
    """
    Функция удаления "плохих" элементов словаря.

    Элемент удвляется, если его значение - запрещённый символ.
    """
    items = list(dictionary.items())
    for key, value in items:
        if value == forbidden_symbol:
            del dictionary[key]
    return dictionary


def removing_symbols(text: str, symbols: Union[str, list[str]],
                     new_s: str = '') -> str:
    """Функция удаления заданных символов из текста."""
    if isinstance(symbols, str):
        return text.replace(symbols, new_s)
    else:
        for symbol in symbols:
            text = text.replace(symbol, new_s)
        return text


def removing_symbols_from_dictionary(dictionary: dict[str, str],
                                     symbols: Union[str, list[str]],
                                     new_s: str = '') -> dict[str, str]:
    """Функция удаления заданных символов из словаря."""
    for key, value in dictionary.items():
        dictionary[key] = removing_symbols(value, symbols, new_s)
    return dictionary


def remove_chars_from_start(text: str, chars: str) -> str:
    """Функция удаления заданных символов из начала строки."""
    # регулярное выражение для поиска определённых символов
    # в начале строки
    pattern = f'^[{re.escape(chars)}]+'
    return re.sub(pattern, '', text)


def remove_chars_from_start_for_dict(dictionary: dict[str, str],
                                     chars: str) -> dict[str, str]:
    """Функция удаления заданых символов в начале строки."""
    for key, value in dictionary.items():
        dictionary[key] = remove_chars_from_start(value, chars)
    return dictionary


def fix_hyphenated_word(text: str) -> str:
    """Функция исправления переноса части слова."""
    return re.sub(r'(?<=\w)-\s+(?=\w)', '', text)


def fix_hyphenated(dictionary: dict[str, str]) -> dict[str, str]:
    """Функция исправления переносов в тексте словаря."""
    for key, value in dictionary.items():
        dictionary[key] = fix_hyphenated_word(value)
    return dictionary


def removing_back_slash(text: str) -> str:
    """Функция удаления кавычек в тексте."""
    return re.sub(r'"', '', text)


def removing_back_slash_dict(dictionary: dict[str, str]) -> dict[str, str]:
    """Функция удаления обратного слеша и кавычек в словаре."""
    for key, value in dictionary.items():
        dictionary[key] = removing_back_slash(value)
    return dictionary


def is_headline(row: str, search_terms: list[str]) -> bool:
    """Функция определения наличия термина из списка в строке."""
    pattern = '|'.join(map(re.escape, search_terms))
    if re.search(pattern, row):
        return True
    else:
        return False


def removing_terms_from_list(text: str, terms: list[str]) -> str:
    """Функция удаления терминов из указанного списка."""
    # Создаем регулярное выражение, которое найдет любые слова из списка
    # Используем r'\b' для того, чтобы удалять только полные слова
    pattern = r'(' + '|'.join(map(re.escape, terms)) + r')'
    # Заменяем все вхождения слов на пустую строку
    result = re.sub(pattern, '', text, flags=re.IGNORECASE)
    # Удаляем лишние пробелы, которые могли образоваться после удаления слов
    result = re.sub(r'\s+', ' ', result).strip()
    return result


def insert_space(text: str) -> str:
    """Функция вставления пробела перед ; если она вплотную к правой буква."""
    return re.sub(r'(?<=\w);(?=\w)', '; ', text)


def docx2json(docx_file: str, json_file: str) -> bool:
    """Функция для прохождение по содержанию файла в порядке перчисления."""
    # Открываем документ
    doc = Document(docx_file)
    data = dict()
    key = ''
    if LEAFLET_INSERT not in doc.paragraphs[1].text.strip().lower():
        # Итерируем по элементам тела документа
        for child in doc.element.body:
            temp_2 = False
            if isinstance(child, CT_P):  # Это абзац
                print('Вход в абзац!', file=log_file)
                paragraph = Paragraph(child, doc)
                if (paragraph.alignment == WD_ALIGN_PARAGRAPH.RIGHT
                   and paragraph.text.strip().isdigit()):
                    continue
                # print('paragraph.style =', paragraph.style)
                print('paragraph.text =', paragraph.text, file=log_file)
                for run in paragraph.runs:
                    print('run.text =', run.text, file=log_file)
                    run.text = run.text.replace('КОД АТХ', 'Код ATX')
                    if (availability(run.text, FORBIDDEN_WORDS) and run.bold):
                        # print('следующая итерация')
                        continue
                    if (run.bold and (not run.italic)
                       and (not run.underline)
                       and ('•' not in run.text)
                       and not temp_2
                       and run.text != ''
                       and not run.text.isspace()
                       and run.text not in string.punctuation
                       and is_first_letter_uppercase(run.text.strip())
                       and not run.text.strip().isupper()
                       and run.text.strip()
                       and not is_headline(run.text, HEADLINES)):
                        print('текст ключа =', run.text, file=log_file)
                        print('начало ключа', file=log_file)
                        if (find_index(run.text.strip(), ABBREVIATIONS) != -1
                           and find_index(run.text.strip(), ABBREVIATIONS) != 0):
                            print('название фирмы с не нулевым индексом =', run.text, file=log_file)
                            index = find_index(run.text.strip(), ABBREVIATIONS)
                            print('index =', index, file=log_file)
                            key = run.text.strip().replace(':', '').replace('.', '')[0:index-1]
                            print('key c индексом =', run.text, file=log_file)
                            data[key] = run.text.strip().replace(':', '').replace('.', '')[index:] + ' '
                        elif find_index(run.text.strip(), ABBREVIATIONS) == 0:
                            print('название фирмы с нулевым индексом =', run.text, file=log_file)
                            print('с ключём: ', key, file=log_file)
                            print('добавляемое значение ', run.text.strip().replace(':', '').replace('.', ''), file=log_file)
                            print('data[key] до добавления', data[key], file=log_file)
                            data[key] += run.text.strip().replace(':', '').replace('.', '') + ' '
                            print('data[key] после добавления', data[key], file=log_file)
                        else:
                            key = run.text.strip().replace(':', '').replace('.', '')
                            print('key без индексом =', run.text, file=log_file)
                            data[key] = ''
                        print('key =', key, file=log_file)
                        temp_2 = True
                    elif (run.bold and (not run.italic)
                          and (not run.underline)
                          and ('•' not in run.text)
                          and temp_2
                          and run.text != ''
                          and not run.text.isspace()
                          and run.text not in string.punctuation
                          and is_first_letter_uppercase(run.text.strip())
                          and not run.text.strip().isupper()
                          and not is_headline(run.text, HEADLINES)):
                        print('текст ключа =', run.text, file=log_file)
                        print('продолжение ключа', file=log_file)
                        if (find_index(run.text.strip(), ABBREVIATIONS) != -1
                           and find_index(run.text.strip(), ABBREVIATIONS) != 0):
                            print('название фирмы с не нулевым индексом =', run.text, file=log_file)
                            index = find_index(run.text.strip(), ABBREVIATIONS)
                            print('index =', index, file=log_file)
                            key += ' ' + run.text.strip().replace(':', '').replace('.', '')[0:index-1]
                            print('key c индексом =', run.text, file=log_file)
                            data[key] = run.text.strip().replace(':', '').replace('.', '')[index:] + ' '
                        elif find_index(run.text.strip(), ABBREVIATIONS) == 0:
                            print('название фирмы с нулевым индексом =', run.text, file=log_file)
                            print('с ключём: ', key, file=log_file)
                            print('добавляемое значение ', run.text.strip().replace(':', '').replace('.', ''), file=log_file)
                            print('data[key] до добавления', data[key], file=log_file)
                            data[key] += run.text.strip().replace(':', '').replace('.', '') + ' '
                            print('data[key] после добавления', data[key], file=log_file)
                        else:
                            key += ' ' + run.text.strip().replace(':', '').replace('.', '') 
                            print('key без индексом =', run.text, file=log_file)
                            #data[key] = run.text.strip().replace(':', '').replace('.', '')[index:]
                            #key += ' ' + run.text.strip().replace(':', '').replace('.', '')
                            data[key] = ''
                        print('key =', key)
                    elif key and temp_2:
                        print('key =', key, file=log_file)
                        print('текст содержания =', run.text, file=log_file)
                        data[key] = run.text.strip() + " "
                        # print('текст содержания =', run.text)
                        temp_2 = False
                    elif key and not temp_2:
                        print('key =', key, file=log_file)
                        print('текст содержания =', run.text, file=log_file)
                        # print('data.keys() =', data.keys())
                        data[key] += run.text.strip() + " "
                        # print('текст содержания =', run.text)
                    else:
                        # print('текст никуда не попал =', run.text)
                        continue

            elif isinstance(child, CT_Tbl):  # Это таблица
                print('Вход в таблицу!', file=log_file)
                print('key =', key, file=log_file)
                table = Table(child, doc)
                first_row = table.rows[0]
                cell = first_row.cells[0]
                print('cell.text =', cell.text, file=log_file)
                if (re.search(ORGAN, cell.text, flags=re.IGNORECASE)
                   and re.search(SYSTEM, cell.text, flags=re.IGNORECASE)):
                    print('таблица ПД', file=log_file)
                    side_effects_data = ''
                    for row in table.rows[2:]:
                        for cell in row.cells:
                            # Параграф с индексом 0 в ячейке
                            if cell.text.strip() == '-':
                                continue
                            table_content = ''
                            if cell.tables:
                                for cell_table in cell.tables:
                                    table_content += get_table_content(
                                        cell_table)
                            cleared_text = table_content + cell.text.strip()
                            cleared_text = clear_text(cleared_text)
                            side_effects_data += cleared_text.strip() + ' '
                    data[KEY_SIDE_EFFECTS] += side_effects_data.strip()
                else:
                    for row in table.rows:
                        temp = False
                        for cell in row.cells:
                            # cell_text = ""
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    print('run.text =', run.text, file=log_file)
                                    print('temp =', temp, file=log_file)
                                    if (availability(run.text, FORBIDDEN_WORDS) and run.bold):
                                        # print('следующая итерация')
                                        continue
                                    if (run.bold and (not run.italic)
                                       and (not run.underline)
                                       and ('•' not in run.text)
                                       and not temp
                                       and run.text != ''
                                       and not run.text.isspace()
                                       #and run.text not in string.punctuation
                                       #and is_first_letter_uppercase(run.text.strip())
                                       #and not run.text.strip().isupper()
                                       #and not availability(run.text.strip(), ABBREVIATIONS)
                                       ):
                                        key = run.text.strip().replace(':', '').replace('.', '')
                                        temp = True
                                        data[key] = ''
                                        print('temp =', temp, file=log_file)
                                        print('1 текст ключа =', run.text, file=log_file)
                                    elif (run.bold and (not run.italic)
                                          and (not run.underline)
                                          and ('•' not in run.text)
                                          and temp
                                          and run.text != ''
                                          and not run.text.isspace()
                                          #and run.text not in string.punctuation
                                          #and is_first_letter_uppercase(run.text.strip())
                                          #and not run.text.strip().isupper()
                                          #and not availability(run.text.strip(), ABBREVIATIONS)
                                          ):
                                        key += ' ' + run.text.strip().replace(':', '').replace('.', '')
                                        data[key] = ''
                                        print('temp =', temp, file=log_file)
                                        print('2 текст ключа =', run.text, file=log_file)
                                    elif key and temp:
                                        temp = False
                                        data[key] = ' ' + run.text.strip() + ' '
                                        print('1 текст содержания =', run.text, file=log_file)
                                    elif key and not temp:
                                        data[key] += ' ' + run.text.strip() + ' '
                                        print('2 текст содержания =', run.text, file=log_file)
                                    else:
                                        # print('текст никуда не попал =', run.text)
                                        temp = False
                                        print('temp в else=', temp, file=log_file)
                                        continue
                                    # if run.bold:
                                    #    cell_text += f"**{run.text}** "
                                    # else:
                                    #    cell_text += run.text + " "
                                print('temp после цикла =', temp, file=log_file)
        print('data.keys =', data.keys(), file=log_file)
        if '' in data.keys():
            print('есть пустой ключ!!!')
            print('есть пустой ключ!!!', file=log_file)
        data = delete_empty_items(data)
        data = removing_forbidden_symbol(data, FORBIDDEN_SYMBOL_R)
        data = removing_symbols_from_dictionary(data, SHY, HYPHEN_SIGN)
        data = remove_chars_from_start_for_dict(data, ': ')
        data = fix_hyphenated(data)
        data = removing_back_slash_dict(data)
        data = removing_extra_spaces_value(data)
        data = delete_empty_items(data)
        if SIDEEFFECT in data:
            data[SIDEEFFECT] = removing_terms_from_list(data[SIDEEFFECT],
                                                        organ_systems)
            data[SIDEEFFECT] = removing_terms_from_list(data[SIDEEFFECT],
                                                        frequency_list)
            data[SIDEEFFECT] = removing_terms_from_list(
                                        data[SIDEEFFECT],
                                        column_names)
            data[SIDEEFFECT] = remove_chars_from_start(
                                        data[SIDEEFFECT], ': ')
            data[SIDEEFFECT] = remove_chars_from_start(
                                        data[SIDEEFFECT], '* ')
            data[SIDEEFFECT] = removing_extra_spaces(
                                        data[SIDEEFFECT])
            data[SIDEEFFECT] = removing_spaces_after_point(
                                        data[SIDEEFFECT])
            data[SIDEEFFECT] = data[SIDEEFFECT].replace('.:', '; ')
            data[SIDEEFFECT] = data[SIDEEFFECT].replace(';:', ';')
            data[SIDEEFFECT] = data[SIDEEFFECT].replace('::', ':')
            data[SIDEEFFECT] = data[SIDEEFFECT].replace('*§', '')
            data[SIDEEFFECT] = data[SIDEEFFECT].replace(' *;', '')
            data[SIDEEFFECT] = data[SIDEEFFECT].replace(' *,', '')
            data[SIDEEFFECT] = data[SIDEEFFECT].replace('*', '')
            data[SIDEEFFECT] = data[SIDEEFFECT].replace(' §', '')
            data[SIDEEFFECT] = data[SIDEEFFECT].replace(' : ', '')
            data[SIDEEFFECT] = data[SIDEEFFECT].replace(' та ', '')
            data[SIDEEFFECT] = data[SIDEEFFECT].replace(' та:', ' ')
            data[SIDEEFFECT] = data[SIDEEFFECT].replace(' ты ', '')
            data[SIDEEFFECT] = data[SIDEEFFECT].replace('(та ', '(')
            data[SIDEEFFECT] = data[SIDEEFFECT].replace(' ификацией',
                                                        ' классификацией')
            data[SIDEEFFECT] = data[SIDEEFFECT].replace(
                                        'Нарушения системы,,,,',
                                        '')
            data[SIDEEFFECT] = data[SIDEEFFECT].replace(
                                        ' слуха и лабиринтные органа нарушения',
                                        '')
            data[SIDEEFFECT] = data[SIDEEFFECT].lstrip(', ')
            data[SIDEEFFECT] = data[SIDEEFFECT].lstrip('та ')
            data[SIDEEFFECT] = insert_space(data[SIDEEFFECT])
            data[SIDEEFFECT] = removing_extra_spaces(data[SIDEEFFECT])
        elif ADVERSE_REACTIONS in data:
            data[ADVERSE_REACTIONS] = removing_terms_from_list(
                                        data[ADVERSE_REACTIONS],
                                        organ_systems)
            data[ADVERSE_REACTIONS] = removing_terms_from_list(
                                        data[ADVERSE_REACTIONS],
                                        frequency_list)
            data[ADVERSE_REACTIONS] = removing_terms_from_list(
                                        data[ADVERSE_REACTIONS],
                                        column_names)
            data[ADVERSE_REACTIONS] = remove_chars_from_start(
                                        data[ADVERSE_REACTIONS], ': ')
            data[ADVERSE_REACTIONS] = remove_chars_from_start(
                                        data[ADVERSE_REACTIONS], '* ')
            data[ADVERSE_REACTIONS] = removing_spaces_after_point(
                                        data[ADVERSE_REACTIONS])
            data[ADVERSE_REACTIONS] = removing_extra_spaces(
                                        data[ADVERSE_REACTIONS])
            data[ADVERSE_REACTIONS] = data[ADVERSE_REACTIONS].replace('.:',
                                                                      ';')
            data[ADVERSE_REACTIONS] = data[ADVERSE_REACTIONS].replace(';:',
                                                                      ';')
            data[ADVERSE_REACTIONS] = data[ADVERSE_REACTIONS].replace('::',
                                                                      ':')
            data[ADVERSE_REACTIONS] = data[ADVERSE_REACTIONS].replace('*§',
                                                                      '')
            data[ADVERSE_REACTIONS] = data[ADVERSE_REACTIONS].replace(' *;',
                                                                      '')
            data[ADVERSE_REACTIONS] = data[ADVERSE_REACTIONS].replace(' *,',
                                                                      '')
            data[ADVERSE_REACTIONS] = data[ADVERSE_REACTIONS].replace('*',
                                                                      '')
            data[ADVERSE_REACTIONS] = data[ADVERSE_REACTIONS].replace(' §',
                                                                      '')
            data[ADVERSE_REACTIONS] = data[ADVERSE_REACTIONS].replace(' : ',
                                                                      '')
            data[ADVERSE_REACTIONS] = data[ADVERSE_REACTIONS].replace(' та ',
                                                                      '')
            data[ADVERSE_REACTIONS] = data[ADVERSE_REACTIONS].replace(' та: ',
                                                                      '')
            data[ADVERSE_REACTIONS] = data[ADVERSE_REACTIONS].replace(' ты ',
                                                                      '')
            data[ADVERSE_REACTIONS] = data[ADVERSE_REACTIONS].replace('(та ',
                                                                      '(')
            data[ADVERSE_REACTIONS] = data[ADVERSE_REACTIONS].replace(
                                        ' ификацией',
                                        ' классификацией')
            data[ADVERSE_REACTIONS] = data[ADVERSE_REACTIONS].replace(
                                        'Нарушения системы,,,,',
                                        '')
            data[ADVERSE_REACTIONS] = data[ADVERSE_REACTIONS].replace(
                                        ' слуха и лабиринтные органа нарушения',
                                        '')
            data[ADVERSE_REACTIONS] = data[ADVERSE_REACTIONS].lstrip(', ')
            data[ADVERSE_REACTIONS] = data[ADVERSE_REACTIONS].lstrip('та ')
            data[ADVERSE_REACTIONS] = insert_space(data[ADVERSE_REACTIONS])
            data[ADVERSE_REACTIONS] = removing_extra_spaces(
                                        data[ADVERSE_REACTIONS])
        with open(json_file, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=4)
        return True
    else:
        return False


def main() -> None:
    """Тело программы."""
    files: list[str] = os.listdir(DIR_DOCX_INSTRUCTIONS)
    if not os.path.exists(DIR_JSON_INSTRUCTIONS):
        os.mkdir(DIR_JSON_INSTRUCTIONS)
    # docx_file = 'Апиксабан.docx'
    # json_file = 'Апиксабан.json'
    for docx_file in files:
        json_file = Path(docx_file).stem + '.json'
        docx_file = os.path.join(DIR_DOCX_INSTRUCTIONS, docx_file)
        json_file = os.path.join(DIR_JSON_INSTRUCTIONS, json_file)
        if docx2json(docx_file, json_file):
            print(f'Файл {docx_file} успешно конвертирован в {json_file}')
        else:
            print('Файл {docx_file} является не инструкцией,'
                  ' а листом-вкладышем!')
        # break
    print('Работы программы успешно завершина!')


if __name__ == '__main__':
    main()


log_file.close()
