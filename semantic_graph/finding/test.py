import pymorphy3
import re
from docx import Document
from docx.shared import RGBColor
import json

morph = pymorphy3.MorphAnalyzer()

# Чтение текстового файла
with open('one_corpus\\corpus_modified.txt', 'r', encoding='utf-8') as file:
    text = file.read()

# Открытие json файла
with open('Dictionary.json', 'r', encoding='utf-8-sig') as file:
    data = json.load(file)

# Разделение текста на слова и пробелы
words = re.split(r'(\s+)', text)
new_words = []

# Создание документа
doc = Document()
paragraph = doc.add_paragraph()

# Функция для замены слова на ключ с сохранением прилегающих символов
def replace_word_with_key(word, key, prefix, suffix):
    print(f"{prefix + word + suffix}  => {prefix + key + suffix}")
    return prefix + key + suffix

for word in words:
    # Проверка, является ли слово буквенным
    if re.match(r'\w+', word):
        # Отделение прилегающих символов
        prefix = re.match(r'^[^\w]+', word)
        suffix = re.search(r'[^\w]+$', word)
        
        if prefix:
            prefix = prefix.group()
            word = word[len(prefix):]
        else:
            prefix = ''
        
        if suffix:
            suffix = suffix.group()
            word = word[:-len(suffix)]
        else:
            suffix = ''

        
        run = paragraph.add_run(prefix)
        run = paragraph.add_run(word)

        # Распознавание глаголов и существительных, обозначающих действие
        lemma = morph.parse(word.lower())[0].normal_form
        pos = morph.parse(word)[0].tag.POS
        if pos == 'NOUN' and lemma.endswith('ие'):
            run.font.color.rgb = RGBColor(0, 255, 255)
        elif pos in {'VERB', 'INFN'}:
            run.font.color.rgb = RGBColor(0, 0, 255)

        # Распознанные синонимы
        synonyms_dict = data["Synonyms"]
        for key, value in synonyms_dict.items():
            if word.lower() == key:
                run.font.color.rgb = RGBColor(0, 255, 0)
                break

        # Распознанные сокращения
        abbreviations_list = data["Abbreviations"]
        for abbreviation in abbreviations_list:
            if  word == abbreviation :
                run.font.color.rgb = RGBColor(255, 0, 0)
                break

        # Распознавание побочных эффектов
        side_effects_list = data["SideEffects"]
        for side_effect in side_effects_list:
            lemma_side_effect = morph.parse(side_effect.lower())[0].normal_form
            if lemma_side_effect == lemma:
                run.font.color.rgb = RGBColor(255, 255, 0)
                print(f"{lemma_side_effect} = {lemma}")
                break

        run = paragraph.add_run(suffix + ' ')
        
# Сохранение документа
doc.save('text_analysis_after_1.docx')